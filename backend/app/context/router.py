"""CRUD-API für context_nodes.

Minimalimplementierung für KS-Phase-1 und -2-Tests.
Sichtbarkeitsfilter werden in KS-Phase-3 um group_memberships-Prüfung erweitert.
"""

import io
import logging
import os
import re
from uuid import UUID

from fastapi import APIRouter, Depends, Form, HTTPException, Query, UploadFile
from sqlalchemy import and_, or_, select, text
from sqlalchemy.ext.asyncio import AsyncSession

import sqlalchemy as sa

from sqlalchemy.exc import IntegrityError

from app.auth.dependencies import get_current_user, require_any_role
from app.auth.jwt import JwtPayload
from app.context.schemas import (
    ContextAnchorCreate,
    ContextAnchorRead,
    ContextEdgeRead,
    ContextEdgeCreate,
    ContextNodeCreate,
    ContextNodeRead,
    ContextNodeUpdate,
    NeighborhoodResponse,
    ArchivedReferenceRead,
    ContextNodeCopyRequest,
    ChatContextNodeAdd,
    ChatContextNodeRead,
    ContextSearchRequest,
    ContextSearchResult,
    CurriculumRead,
    CurriculumDraft,
    CurriculumDraftConfirmed,
    CurriculumCreate,
)
from app.context.embedding import enqueue_embedding_job
from app.context.taxonomy import validate_content_type
from app.context.retrieval import VALID_SCOPE_ANCHOR_TYPES
from app.preferences.service import get_preferences
from app.db.models import (
    Assistant,
    AssistantContextAnchor,
    ChatContextNode,
    ContextEdge,
    ContextNode,
    Conversation,
    Group,
    GroupMembership,
    Subject,
)
from app.db.session import get_db

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/context", tags=["context"])

_TEACHER_OR_ADMIN = require_any_role(["teacher", "admin"])


async def _check_anchor_permission(
    assistant_id: int,
    db: AsyncSession,
    current_user: JwtPayload,
) -> Assistant:
    """Laedt den Assistenten und prueft Schreibrecht (Eigentuemer oder Admin)."""
    assistant = await db.get(Assistant, assistant_id)
    if assistant is None:
        raise HTTPException(status_code=404, detail="Assistent nicht gefunden")
    if "admin" not in current_user.roles and assistant.created_by != current_user.sub:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")
    return assistant


def _check_write_permission(node: ContextNode, user: JwtPayload) -> None:
    """403 wenn weder Admin noch owner_pseudonym."""
    if "admin" not in user.roles and node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")


def _visibility_filter(query, user: JwtPayload, status_override: str | None = None):
    """Sichtbarkeitsfilter; status_override überschreibt den active-Default."""
    q = query.where(
        or_(
            ContextNode.read_scope.in_(["global", "school", "subject", "group"]),
            ContextNode.owner_pseudonym == user.sub,
        ),
    )
    if status_override is not None:
        q = q.where(ContextNode.status == status_override)
    else:
        q = q.where(ContextNode.status == "active")
    return q


# ── GET /api/context/nodes ────────────────────────────────────────────────────


@router.get("/nodes", response_model=list[ContextNodeRead])
async def list_nodes(
    q: str | None = Query(default=None),
    category: str | None = Query(default=None),
    content_type: list[str] | None = Query(default=None),
    status: str | None = Query(default=None),
    subject_slug: str | None = Query(default=None),
    group_id: int | None = Query(default=None),
    grade: int | None = Query(default=None, ge=1, le=13, description="Jahrgangsstufe"),
    owner: str | None = Query(default=None),
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    query = select(ContextNode)
    query = _visibility_filter(query, user, status_override=status)

    # subject_slug: Knoten deren Scope-Gruppe zu diesem Fach gehört,
    # plus schulweite/globale knowledge-Knoten mit passendem Fach oder fächerübergreifend
    if subject_slug:
        subquery_group_ids = (
            sa.select(Group.id)
            .join(Subject, Subject.id == Group.subject_id)
            .where(Subject.slug == subject_slug)
            .scalar_subquery()
        )
        # subject_id der Ziel-Slug nachschlagen
        subject_id_subq = (
            sa.select(Subject.id)
            .where(Subject.slug == subject_slug)
            .scalar_subquery()
        )
        query = query.where(
            or_(
                ContextNode.read_scope_group_id.in_(subquery_group_ids),
                and_(
                    ContextNode.read_scope.in_(["global", "school"]),
                    ContextNode.category == "knowledge",
                    or_(
                        ContextNode.subject_id == subject_id_subq,
                        ContextNode.subject_id.is_(None),
                    ),
                ),
            )
        )

    # group_id: nur Knoten mit exakt dieser read_scope_group_id
    if group_id is not None:
        query = query.where(ContextNode.read_scope_group_id == group_id)

    # grade: Jahrgangsstufen-Filter
    if grade is not None:
        query = query.where(
            or_(
                ContextNode.min_grade.is_(None),  # keine Stufenangabe = für alle
                and_(
                    ContextNode.min_grade <= grade,
                    ContextNode.max_grade >= grade,
                ),
            )
        )

    # owner=me: nur eigene Knoten
    if owner is not None:
        if owner != "me":
            raise HTTPException(status_code=400, detail="owner muss 'me' sein")
        query = query.where(ContextNode.owner_pseudonym == user.sub)

    if q:
        query = query.where(ContextNode.title.ilike(f"%{q}%"))
    if category:
        query = query.where(ContextNode.category == category)
    if content_type:
        query = query.where(ContextNode.content_type.in_(content_type))

    result = await db.execute(query.order_by(ContextNode.created_at.desc()))
    return result.scalars().all()


# ── GET /api/context/nodes/{id}/neighborhood ────────────────────────────────


@router.get("/nodes/{node_id}/neighborhood", response_model=NeighborhoodResponse)
async def get_neighborhood(
    node_id: UUID,
    depth: int = Query(default=2, ge=1, le=3),
    relation: list[str] | None = Query(default=None),
    category: list[str] | None = Query(default=None),
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    # Startknoten laden und prüfen
    node = await db.get(ContextNode, node_id)
    if node is None or node.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")

    # Recursive CTE für bidirektionale Traversierung
    neighborhood_cte = text("""
        WITH RECURSIVE nb AS (
            SELECT id, 0 AS depth
              FROM context_nodes WHERE id = :node_id
            UNION
            SELECT
                CASE WHEN e.from_node_id = nb.id THEN e.to_node_id
                     ELSE e.from_node_id END,
                nb.depth + 1
              FROM nb
              JOIN context_edges e
                ON e.from_node_id = nb.id OR e.to_node_id = nb.id
              JOIN context_nodes n
                ON n.id = CASE WHEN e.from_node_id = nb.id
                               THEN e.to_node_id ELSE e.from_node_id END
             WHERE nb.depth < :depth
               AND n.status = 'active'
        )
        SELECT DISTINCT id FROM nb
    """)
    result = await db.execute(neighborhood_cte, {"node_id": str(node_id), "depth": depth})
    neighbor_ids = [row[0] for row in result.fetchall()]

    # Knoten laden + Sichtbarkeitsfilter anwenden
    nodes_query = select(ContextNode).where(
        ContextNode.id.in_(neighbor_ids),
        ContextNode.status == "active",
        or_(
            ContextNode.read_scope.in_(["global", "school", "subject", "group"]),
            ContextNode.owner_pseudonym == user.sub,
        ),
    )
    if category:
        nodes_query = nodes_query.where(ContextNode.category.in_(category))
    nodes = (await db.execute(nodes_query)).scalars().all()
    visible_ids = {n.id for n in nodes}

    # Kanten zwischen sichtbaren Knoten laden
    edges_query = select(ContextEdge).where(
        ContextEdge.from_node_id.in_(visible_ids),
        ContextEdge.to_node_id.in_(visible_ids),
    )
    if relation:
        edges_query = edges_query.where(ContextEdge.relation.in_(relation))
    edges = (await db.execute(edges_query)).scalars().all()

    return NeighborhoodResponse(nodes=nodes, edges=edges)


# ── GET /api/context/nodes/{id}/archived-references ────────────────────────


@router.get(
    "/nodes/{node_id}/archived-references",
    response_model=list[ArchivedReferenceRead],
)
async def get_archived_references(
    node_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    # Startknoten laden
    node = await db.get(ContextNode, node_id)
    if node is None or node.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")

    sql = text("""
        SELECT
            n.id,
            n.title,
            n.category,
            n.content_type,
            n.archived_at,
            e.relation,
            (
                SELECT s.id FROM context_edges se
                  JOIN context_nodes s ON s.id = se.from_node_id
                 WHERE se.to_node_id = n.id
                   AND se.relation = 'supersedes'
                   AND s.status = 'active'
                 LIMIT 1
            ) AS suggested_successor_id
          FROM context_edges e
          JOIN context_nodes n ON n.id = e.to_node_id
         WHERE e.from_node_id = :node_id
           AND n.status = 'archived'
    """)
    result = await db.execute(sql, {"node_id": str(node_id)})
    rows = result.mappings().all()
    return [ArchivedReferenceRead(**dict(row)) for row in rows]


# ── POST /api/context/nodes/{id}/copy ───────────────────────────────────────


@router.post(
    "/nodes/{node_id}/copy",
    response_model=ContextNodeRead,
    status_code=201,
)
async def copy_node(
    node_id: UUID,
    payload: ContextNodeCopyRequest,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    # Quellknoten laden
    source = await db.get(ContextNode, node_id)
    if source is None or source.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")

    # Neuen Knoten anlegen
    new_node = ContextNode(
        category=source.category,
        content_type=source.content_type,
        title=source.title,
        content=source.content,
        metadata_=source.metadata_,
        owner_pseudonym=user.sub,
        read_scope=source.read_scope,
        write_scope=source.write_scope,
        read_scope_group_id=payload.read_scope_group_id or source.read_scope_group_id,
        write_scope_group_id=payload.write_scope_group_id or source.write_scope_group_id,
        valid_until=payload.valid_until,
        schuljahr=payload.schuljahr or source.schuljahr,
        status="active",
    )
    db.add(new_node)
    await db.commit()
    await db.refresh(new_node)
    await enqueue_embedding_job(new_node.id, db)
    return new_node


# ── GET /api/context/nodes/{id} ─────────────────────────────────────────────────

@router.get("/nodes/{node_id}", response_model=ContextNodeRead)
async def get_node(
    node_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    node = await db.get(ContextNode, node_id)
    if node is None or node.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")
    if node.read_scope == "private" and node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")
    return node


# ── POST /api/context/nodes ────────────────────────────────────────────────────

@router.post("/nodes", response_model=ContextNodeRead, status_code=201)
async def create_node(
    payload: ContextNodeCreate,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    try:
        validate_content_type(payload.category, payload.content_type)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))

    node = ContextNode(
        category=payload.category,
        content_type=payload.content_type,
        title=payload.title,
        content=payload.content,
        metadata_=payload.metadata_,
        owner_pseudonym=payload.owner_pseudonym or user.sub,
        read_scope=payload.read_scope,
        write_scope=payload.write_scope,
        read_scope_group_id=payload.read_scope_group_id,
        write_scope_group_id=payload.write_scope_group_id,
        assistant_id=payload.assistant_id,
        subject_id=payload.subject_id,
        min_grade=payload.min_grade,
        max_grade=payload.max_grade,
        valid_until=payload.valid_until,
        schuljahr=payload.schuljahr,
    )
    db.add(node)
    await db.commit()
    await db.refresh(node)

    # Embedding-Job
    await enqueue_embedding_job(node.id, db)

    return node


# ── PATCH /api/context/nodes/{id} ─────────────────────────────────────────────

@router.patch("/nodes/{node_id}", response_model=ContextNodeRead)
async def update_node(
    node_id: UUID,
    payload: ContextNodeUpdate,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    node = await db.get(ContextNode, node_id)
    if node is None or node.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")
    _check_write_permission(node, user)

    update_data = payload.model_dump(exclude_unset=True, by_alias=False)
    for field, value in update_data.items():
        # metadata_ → DB-Spalte 'metadata'
        attr = field if field != "metadata_" else "metadata_"
        setattr(node, attr, value)

    await db.commit()
    await db.refresh(node)
    return node


# ── DELETE /api/context/nodes/{id} ────────────────────────────────────────────

@router.delete("/nodes/{node_id}", status_code=204)
async def delete_node(
    node_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    node = await db.get(ContextNode, node_id)
    if node is None or node.status == "deleted":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")
    _check_write_permission(node, user)
    await db.delete(node)
    await db.commit()


# ── Context Anchors ────────────────────────────────────────────────────────

@router.get(
    "/assistants/{assistant_id}/anchors",
    response_model=list[ContextAnchorRead],
)
async def list_context_anchors(
    assistant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: JwtPayload = Depends(get_current_user),
) -> list[ContextAnchorRead]:
    await _check_anchor_permission(assistant_id, db, current_user)

    result = await db.execute(
        sa.select(
            AssistantContextAnchor.assistant_id,
            AssistantContextAnchor.node_id,
            AssistantContextAnchor.role,
            ContextNode.title.label("node_title"),
            ContextNode.content_type.label("node_content_type"),
        )
        .join(ContextNode, ContextNode.id == AssistantContextAnchor.node_id)
        .where(AssistantContextAnchor.assistant_id == assistant_id)
    )
    rows = result.mappings().all()
    return [ContextAnchorRead(**dict(row)) for row in rows]


@router.post(
    "/assistants/{assistant_id}/anchors",
    response_model=ContextAnchorRead,
    status_code=201,
)
async def add_context_anchor(
    assistant_id: int,
    body: ContextAnchorCreate,
    db: AsyncSession = Depends(get_db),
    current_user: JwtPayload = Depends(get_current_user),
) -> ContextAnchorRead:
    await _check_anchor_permission(assistant_id, db, current_user)

    # Knoten laden und validieren
    node = await db.get(ContextNode, body.node_id)
    if node is None or node.status != "active":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden oder inaktiv")

    # Fuer retrieval_scope: nur strukturell sinnvolle Typen zulassen
    if body.role == "retrieval_scope" and node.content_type not in VALID_SCOPE_ANCHOR_TYPES:
        raise HTTPException(
            status_code=422,
            detail=(
                f"content_type '{node.content_type}' ist kein gueltiger retrieval_scope-Anker. "
                f"Erlaubt: {sorted(VALID_SCOPE_ANCHOR_TYPES)}"
            ),
        )

    anchor = AssistantContextAnchor(
        assistant_id=assistant_id,
        node_id=body.node_id,
        role=body.role,
    )
    db.add(anchor)
    try:
        await db.flush()
    except IntegrityError:
        await db.rollback()
        raise HTTPException(status_code=409, detail="Anker bereits vorhanden")

    await db.commit()
    await db.refresh(node)

    return ContextAnchorRead(
        assistant_id=assistant_id,
        node_id=body.node_id,
        role=body.role,
        node_title=node.title,
        node_content_type=node.content_type,
    )


@router.delete(
    "/assistants/{assistant_id}/anchors/{node_id}/{role}",
    status_code=204,
)
async def remove_context_anchor(
    assistant_id: int,
    node_id: UUID,
    role: str,
    db: AsyncSession = Depends(get_db),
    current_user: JwtPayload = Depends(get_current_user),
) -> None:
    await _check_anchor_permission(assistant_id, db, current_user)

    anchor = await db.get(
        AssistantContextAnchor, (assistant_id, node_id, role)
    )
    if anchor is None:
        raise HTTPException(status_code=404, detail="Anker nicht gefunden")
    await db.delete(anchor)
    await db.commit()


# ── KS-Phase-5 Chat Context Nodes ──────────────────────────────────────────


async def _get_conversation_or_403(
    conversation_id: UUID,
    pseudonym: str,
    db: AsyncSession,
) -> None:
    """404 wenn Konversation nicht existiert, 403 wenn sie nicht dem Nutzer gehört."""
    conv = await db.get(Conversation, conversation_id)
    if conv is None:
        raise HTTPException(status_code=404, detail="Konversation nicht gefunden")
    if conv.pseudonym != pseudonym:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")


@router.get(
    "/conversations/{conversation_id}/nodes",
    response_model=list[ChatContextNodeRead],
)
async def list_chat_context_nodes(
    conversation_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    await _get_conversation_or_403(conversation_id, user.sub, db)

    result = await db.execute(
        sa.select(
            ChatContextNode.node_id,
            ContextNode.category,
            ContextNode.title,
            ContextNode.content_type,
            ChatContextNode.added_at,
        )
        .join(ContextNode, ContextNode.id == ChatContextNode.node_id)
        .where(
            ChatContextNode.chat_id == conversation_id,
            ContextNode.status == "active",
        )
        .order_by(ChatContextNode.added_at)
    )
    rows = result.mappings().all()
    return [ChatContextNodeRead(**dict(row)) for row in rows]


@router.post(
    "/conversations/{conversation_id}/nodes",
    response_model=ChatContextNodeRead,
    status_code=201,
)
async def add_chat_context_node(
    conversation_id: UUID,
    payload: ChatContextNodeAdd,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    await _get_conversation_or_403(conversation_id, user.sub, db)

    # Knoten existiert und ist aktiv?
    node = await db.get(ContextNode, payload.node_id)
    if node is None or node.status != "active":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden oder inaktiv")

    # Sichtbarkeit prüfen (kein privater Fremd-Knoten)
    if node.read_scope == "private" and node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")

    existing = await db.get(ChatContextNode, (conversation_id, payload.node_id))
    if existing is not None:
        return ChatContextNodeRead(
            node_id=node.id,
            category=node.category,
            title=node.title,
            content_type=node.content_type,
            added_at=existing.added_at,
        )

    entry = ChatContextNode(chat_id=conversation_id, node_id=payload.node_id)
    db.add(entry)
    await db.flush()
    await db.commit()

    return ChatContextNodeRead(
        node_id=node.id,
        category=node.category,
        title=node.title,
        content_type=node.content_type,
        added_at=entry.added_at,
    )


@router.delete(
    "/conversations/{conversation_id}/nodes/{node_id}",
    status_code=204,
)
async def remove_chat_context_node(
    conversation_id: UUID,
    node_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    await _get_conversation_or_403(conversation_id, user.sub, db)

    entry = await db.get(ChatContextNode, (conversation_id, node_id))
    if entry is None:
        raise HTTPException(status_code=404, detail="Eintrag nicht gefunden")
    await db.delete(entry)
    await db.commit()


# ── KS-Phase-5 Semantic Search ──────────────────────────────────────────


@router.post("/search", response_model=list[ContextSearchResult])
async def search_context_nodes(
    request: ContextSearchRequest,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(get_current_user),
):
    """Semantische Suche über sichtbare Knoten anhand eines Freitexts."""
    from app.chat.router import _exec_search_context_nodes
    prefs = await get_preferences(db, user.sub)
    try:
        limit = max(5, min(30, int(prefs.get("context_search_limit", 8))))
    except (TypeError, ValueError):
        limit = 8
    results = await _exec_search_context_nodes(request.query, user.sub, db, limit=limit)
    return results


# ── KS-Phase-6 Curriculum Endpoints ──────────────────────────────────────


@router.get("/curricula/{curriculum_id}", response_model=CurriculumRead)
async def get_curriculum(
    curriculum_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(get_current_user),
):
    """Gibt das vollständige Curriculum als verschachteltes Objekt zurück."""
    import os
    
    # Curriculum-Knoten laden
    result = await db.execute(
        sa.select(ContextNode).where(
            ContextNode.id == curriculum_id,
            ContextNode.status == "active",
            ContextNode.content_type == "curriculum",
        )
    )
    curriculum = result.scalar_one_or_none()
    
    if not curriculum:
        raise HTTPException(status_code=404, detail="Curriculum nicht gefunden oder inaktiv")
    
    # Sichtbarkeit prüfen
    if curriculum.read_scope == "private":
        if curriculum.owner_pseudonym != user.sub:
            raise HTTPException(status_code=403, detail="Keine Berechtigung")
    elif curriculum.read_scope == "subject":
        # Schüler dürfen nur mit ENV-Flag zugreifen
        if "student" in user.roles and "teacher" not in user.roles:
            if os.environ.get("CURRICULUM_VISIBLE_TO_STUDENTS", "false").lower() != "true":
                raise HTTPException(status_code=403, detail="Keine Berechtigung")
    
    # Alle Kapitel zu diesem Curriculum laden (via part_of-Kanten)
    result = await db.execute(
        sa.select(ContextNode)
        .join(ContextEdge, ContextEdge.from_node_id == ContextNode.id)
        .where(
            ContextEdge.to_node_id == curriculum_id,
            ContextEdge.relation == "part_of",
            ContextNode.content_type == "kapitel",
            ContextNode.status == "active",
        )
        .order_by(ContextNode.metadata_["reihenfolge"].as_integer())
    )
    kapitel_nodes = result.scalars().all()

    kapitel_list = []
    for kap_node in kapitel_nodes:
        # Alle Lernsequenzen zu diesem Kapitel laden
        result = await db.execute(
            sa.select(ContextNode)
            .join(ContextEdge, ContextEdge.from_node_id == ContextNode.id)
            .where(
                ContextEdge.to_node_id == kap_node.id,
                ContextEdge.relation == "part_of",
                ContextNode.content_type == "lernsequenz",
                ContextNode.status == "active",
            )
            .order_by(ContextNode.metadata_["reihenfolge"].as_integer())
        )
        lernsequenz_nodes = result.scalars().all()

        lernsequenzen_list = []
        for ls_node in lernsequenz_nodes:
            # IK-Referenzen laden
            result = await db.execute(
                sa.text("""
                    SELECT n.id, n.title, e.metadata->>'partiell' as partiell
                    FROM context_nodes n
                    JOIN context_edges e ON e.to_node_id = n.id
                    WHERE e.from_node_id = :ls_id
                      AND e.relation = 'references'
                      AND n.content_type = 'ik_kompetenz'
                      AND n.status = 'active'
                """),
                {"ls_id": str(ls_node.id)},
            )
            ik_refs = [
                {"node_id": str(row.id), "title": row.title, "partiell": row.partiell == "true"}
                for row in result.mappings().all()
            ]

            # PK-Referenzen laden
            result = await db.execute(
                sa.text("""
                    SELECT n.id, n.title
                    FROM context_nodes n
                    JOIN context_edges e ON e.to_node_id = n.id
                    WHERE e.from_node_id = :ls_id
                      AND e.relation = 'develops'
                      AND n.content_type = 'pk_kompetenz'
                      AND n.status = 'active'
                """),
                {"ls_id": str(ls_node.id)},
            )
            pk_refs = [
                {"node_id": str(row.id), "title": row.title}
                for row in result.mappings().all()
            ]

            # Leitperspektive-Referenzen laden
            result = await db.execute(
                sa.text("""
                    SELECT n.id, n.title, n.metadata->>'code' as lp_code
                    FROM context_nodes n
                    JOIN context_edges e ON e.to_node_id = n.id
                    WHERE e.from_node_id = :ls_id
                      AND e.relation = 'references'
                      AND n.content_type = 'leitperspektive'
                      AND n.status = 'active'
                """),
                {"ls_id": str(ls_node.id)},
            )
            leitperspektive_refs = [
                {"node_id": str(row.id), "title": row.title, "lp_code": row.lp_code}
                for row in result.mappings().all()
            ]
            
            lernsequenzen_list.append({
                "id": ls_node.id,
                "title": ls_node.title,
                "metadata": ls_node.metadata_ or {},
                "ik_refs": ik_refs,
                "pk_refs": pk_refs,
                "leitperspektive_refs": leitperspektive_refs,
            })
        
        kapitel_list.append({
            "id": kap_node.id,
            "title": kap_node.title,
            "metadata": kap_node.metadata_ or {},
            "lernsequenzen": lernsequenzen_list,
        })
    
    # Prüfe ob User editieren darf
    can_edit = False
    if "admin" in user.roles:
        can_edit = True
    elif curriculum.write_scope_group_id:
        result = await db.execute(
            sa.select(1).where(
                sa.exists().where(
                    GroupMembership.group_id == curriculum.write_scope_group_id,
                    GroupMembership.pseudonym == user.sub,
                )
            )
        )
        can_edit = result.scalar_one_or_none() is not None

    return {
        "id": curriculum.id,
        "title": curriculum.title,
        "metadata": curriculum.metadata_ or {},
        "subject_id": curriculum.subject_id,
        "write_scope_group_id": curriculum.write_scope_group_id,
        "kapitel": kapitel_list,
        "can_edit": can_edit,
    }


@router.get("/curricula/by-subject/{subject_id}", response_model=list[ContextNodeRead])
async def list_curricula_by_subject(
    subject_id: int,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(get_current_user),
):
    """Gibt alle Curriculum-Knoten eines Fachs zurück."""
    result = await db.execute(
        sa.select(ContextNode).where(
            ContextNode.content_type == "curriculum",
            ContextNode.subject_id == subject_id,
            ContextNode.status == "active",
            sa.or_(
                ContextNode.read_scope.in_(["global", "school", "subject"]),
                ContextNode.owner_pseudonym == user.sub,
            ),
        ).order_by(ContextNode.metadata_["jahrgangsstufe"])
    )
    curricula = result.scalars().all()
    return curricula


# ── KS-Phase-6 Curriculum Convert (Stufe 1) ─────────────────────────────────

@router.post("/curricula/convert", response_model=CurriculumDraft)
async def convert_curriculum(
    file: UploadFile,
    fachplan_id: str = Form(...),
    fach: str = Form(...),
    jahrgangsstufe: str = Form(...),
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    """Konvertiert PDF/Word-Dokument zu strukturiertem Zwischenformat (Stufe 1).
    
    Nimmt ein PDF- oder Word-Dokument entgegen und extrahiert die Curriculum-Struktur
    mit LLM-Unterstützung. Das Ergebnis ist ein CurriculumDraft für den Prüfungsschritt.
    """
    # Datei-Inhalt lesen
    try:
        content = await file.read()
    except Exception as e:
        logger.error(f"Fehler beim Lesen der Datei: {e}")
        raise HTTPException(status_code=400, detail="Datei konnte nicht gelesen werden")
    
    # Dateityp erkennen
    filename = file.filename or ""
    is_pdf = filename.lower().endswith(".pdf")
    is_word = filename.lower().endswith((".docx", ".doc"))
    
    if not is_pdf and not is_word:
        raise HTTPException(
            status_code=400,
            detail="Nur PDF oder Word-Dokumente (.docx, .doc) werden unterstützt"
        )
    
    # bp_version aus Fachplan-Metadaten lesen
    from app.db.models import ContextNode as ContextNodeModel
    try:
        fachplan_uuid = UUID(fachplan_id)
        fachplan_node = await db.get(ContextNodeModel, fachplan_uuid)
        bp_version = (fachplan_node.metadata_ or {}).get("bp_version", "") if fachplan_node else ""
    except Exception:
        bp_version = ""

    # Schulart aus Konfiguration
    from app.config import settings as _settings
    schulart = _settings.schulart

    # Format erkennen
    is_supported, format_detected, warnings = _detect_format(
        _extract_text_from_pdf(content) if is_pdf else _extract_text_from_word(content)
    )

    if not is_supported:
        return CurriculumDraft(
            unsupported_format=True,
            format_detected=format_detected,
            warnings=warnings,
            data=None
        )

    # LLM-basierte Strukturextraktion
    try:
        if is_pdf:
            draft_data = await _extract_curriculum_via_llm(
                content, fachplan_id, bp_version,
                fach=fach, jahrgangsstufe=jahrgangsstufe, schulart=schulart
            )
        else:
            # Für Word: Serialisierung + LLM-Pfad
            pages = _serialize_word_for_llm(content)
            
            # Temporäre Umwandlung für LLM-Extraktion
            # Da Word-Dokumente typischerweise weniger Seiten haben, können wir alle auf einmal verarbeiten
            all_content_parts = []
            for page in pages:
                if page.tables:
                    all_content_parts.extend(page.tables)
                if page.flow_text.strip():
                    all_content_parts.append(f"--- Fließtext Seite {page.page_number} ---")
                    all_content_parts.append(page.flow_text.strip())
            
            serialized_input = "\n".join(all_content_parts)
            
            # LLM aufrufen für Word-Inhalt
            result = await _call_extraction_llm(serialized_input, 0, _settings)
            chapter_data = result.get("kapitel", {})
            chapter = CurriculumDraftKapitel.model_validate(chapter_data)
            chapter.reihenfolge = 1
            
            draft_data = CurriculumDraftData(
                schule="",
                fach_code=fach[:2].upper() if fach else "??",
                fach=fach or None,
                schulart=schulart,
                jahrgangsstufe=jahrgangsstufe,
                fachplan_id=fachplan_id,
                bp_version=bp_version,
                vorwort=None,
                kapitel=[chapter],
            )
        
        return CurriculumDraft(
            unsupported_format=False,
            format_detected=format_detected,
            warnings=warnings,
            data=draft_data
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"LLM-basierte Strukturextraktion fehlgeschlagen: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"LLM-basierte Strukturextraktion fehlgeschlagen: {str(e)}"
        )


# ── KS-Phase-6 Curriculum Create (Stufe 2) ──────────────────────────────────

@router.post("/curricula", response_model=CurriculumRead, status_code=201)
async def create_curriculum(
    payload: CurriculumDraftConfirmed,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    """Speichert ein bestätigtes Curriculum aus dem Zwischenformat (Stufe 2).
    
    Nimmt das bestätigte Zwischenformat entgegen und erstellt alle Knoten und Kanten.
    """
    from app.context.service import import_curriculum_from_draft
    
    try:
        curriculum_id, stats = await import_curriculum_from_draft(
            db, payload, user.sub
        )
        
        # Log warnings
        for warning in stats.warnings:
            logger.warning(f"Curriculum-Import Warnung: {warning}")
        
        # Lade das erstellte Curriculum für die Rückgabe
        result = await db.execute(
            sa.select(ContextNode).where(ContextNode.id == curriculum_id)
        )
        curriculum = result.scalar_one()
        
        # Baue die verschachtelte Struktur (wie in get_curriculum)
        # Da wir gerade erstellt haben, können wir direkt die IDs verwenden
        kapitel_list = []
        for kap in payload.kapitel:
            # Finde kapitel_id
            kapitel_import_key = f"{payload.fachplan_id}_{payload.jahrgangsstufe}_kapitel_{kap.reihenfolge}"
            kap_result = await db.execute(
                sa.select(ContextNode.id).where(
                    ContextNode.content_type == "kapitel",
                    ContextNode.metadata_["import_key"].astext == kapitel_import_key
                )
            )
            kap_id = kap_result.scalar_one_or_none()
            
            if not kap_id:
                continue
                
            lernsequenzen_list = []
            for ls in kap.lernsequenzen:
                ls_reihenfolge = ls.reihenfolge if ls.reihenfolge is not None else 0
                ls_import_key = f"{kapitel_import_key}_ls_{ls_reihenfolge}"
                ls_result = await db.execute(
                    sa.select(ContextNode).where(
                        ContextNode.content_type == "lernsequenz",
                        ContextNode.metadata_["import_key"].astext == ls_import_key
                    )
                )
                ls_node = ls_result.scalar_one_or_none()
                
                if not ls_node:
                    continue
                
                # IK-Referenzen laden
                ik_result = await db.execute(
                    sa.text("""
                        SELECT n.id, n.title, e.metadata->>'partiell' as partiell
                        FROM context_nodes n
                        JOIN context_edges e ON e.to_node_id = n.id
                        WHERE e.from_node_id = :ls_id
                          AND e.relation = 'references'
                          AND n.content_type = 'ik_kompetenz'
                          AND n.status = 'active'
                    """),
                    {"ls_id": str(ls_node.id)},
                )
                ik_refs = [
                    {"node_id": str(row.id), "title": row.title, "partiell": row.partiell == "true"}
                    for row in ik_result.mappings().all()
                ]

                # PK-Referenzen laden
                pk_result = await db.execute(
                    sa.text("""
                        SELECT n.id, n.title
                        FROM context_nodes n
                        JOIN context_edges e ON e.to_node_id = n.id
                        WHERE e.from_node_id = :ls_id
                          AND e.relation = 'develops'
                          AND n.content_type = 'pk_kompetenz'
                          AND n.status = 'active'
                    """),
                    {"ls_id": str(ls_node.id)},
                )
                pk_refs = [
                    {"node_id": str(row.id), "title": row.title}
                    for row in pk_result.mappings().all()
                ]

                # Leitperspektive-Referenzen laden
                lp_result = await db.execute(
                    sa.text("""
                        SELECT n.id, n.title, n.metadata->>'code' as lp_code
                        FROM context_nodes n
                        JOIN context_edges e ON e.to_node_id = n.id
                        WHERE e.from_node_id = :ls_id
                          AND e.relation = 'references'
                          AND n.content_type = 'leitperspektive'
                          AND n.status = 'active'
                    """),
                    {"ls_id": str(ls_node.id)},
                )
                leitperspektive_refs = [
                    {"node_id": str(row.id), "title": row.title, "lp_code": row.lp_code}
                    for row in lp_result.mappings().all()
                ]
                
                lernsequenzen_list.append({
                    "id": ls_node.id,
                    "title": ls_node.title,
                    "metadata": ls_node.metadata_ or {},
                    "ik_refs": ik_refs,
                    "pk_refs": pk_refs,
                    "leitperspektive_refs": leitperspektive_refs,
                })
            
            kapitel_list.append({
                "id": kap_id,
                "title": kap_node.title if 'kap_node' in locals() else kap.titel,
                "metadata": {},  # Wird unten korrigiert
                "lernsequenzen": lernsequenzen_list,
            })
        
        # Korrigiere Kapitel-Metadata
        for kap in payload.kapitel:
            kapitel_import_key = f"{payload.fachplan_id}_{payload.jahrgangsstufe}_kapitel_{kap.reihenfolge}"
            kap_result = await db.execute(
                sa.select(ContextNode).where(
                    ContextNode.content_type == "kapitel",
                    ContextNode.metadata_["import_key"].astext == kapitel_import_key
                )
            )
            kap_node = kap_result.scalar_one_or_none()
            if kap_node:
                for k in kapitel_list:
                    if str(k["id"]) == str(kap_node.id):
                        k["metadata"] = kap_node.metadata_ or {}
                        k["title"] = kap_node.title
                        break
        
        # Prüfe can_edit
        department_group_id = curriculum.write_scope_group_id
        can_edit = False
        if "admin" in user.roles:
            can_edit = True
        elif department_group_id:
            result = await db.execute(
                sa.select(1).where(
                    sa.exists().where(
                        GroupMembership.group_id == department_group_id,
                        GroupMembership.pseudonym == user.sub,
                    )
                )
            )
            can_edit = result.scalar_one_or_none() is not None
        
        return {
            "id": curriculum.id,
            "title": curriculum.title,
            "metadata": curriculum.metadata_ or {},
            "subject_id": curriculum.subject_id,
            "write_scope_group_id": curriculum.write_scope_group_id,
            "kapitel": kapitel_list,
            "can_edit": can_edit,
        }
        
    except ValueError as e:
        logger.error(f"Validierungsfehler beim Curriculum-Import: {e}")
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.error(f"Fehler beim Curriculum-Import: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── KS-Phase-6 Edge CRUD Endpoints ──────────────────────────────────────


@router.post("/edges", response_model=ContextEdgeRead, status_code=201)
async def create_edge(
    payload: ContextEdgeCreate,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    """Erstellt eine neue Kante zwischen zwei Knoten."""
    from app.db.models import ContextNode
    
    # Prüfe ob beide Knoten existieren
    from_node = await db.get(ContextNode, payload.from_node_id)
    to_node = await db.get(ContextNode, payload.to_node_id)
    
    if not from_node or from_node.status != "active":
        raise HTTPException(status_code=404, detail=f"Startknoten {payload.from_node_id} nicht gefunden")
    if not to_node or to_node.status != "active":
        raise HTTPException(status_code=404, detail=f"Zielknoten {payload.to_node_id} nicht gefunden")
    
    # Prüfe Schreibrecht auf from_node
    if from_node.write_scope == "private" and from_node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Schreibberechtigung auf Startknoten")
    if from_node.write_scope == "subject" or from_node.write_scope == "group":
        if from_node.write_scope_group_id:
            # Prüfe ob User Mitglied der Gruppe ist
            is_member = await db.execute(
                sa.select(1).where(
                    GroupMembership.group_id == from_node.write_scope_group_id,
                    GroupMembership.pseudonym == user.sub,
                )
            )
            if not is_member.scalar_one_or_none():
                raise HTTPException(status_code=403, detail="Keine Schreibberechtigung auf Startknoten")
    
    # Prüfe ob Kante bereits existiert (idempotent)
    existing = await db.execute(
        sa.select(ContextEdge).where(
            ContextEdge.from_node_id == payload.from_node_id,
            ContextEdge.to_node_id == payload.to_node_id,
            ContextEdge.relation == payload.relation,
        )
    )
    existing_edge = existing.scalar_one_or_none()
    if existing_edge:
        return existing_edge
    
    # Kante erstellen
    edge = ContextEdge(
        from_node_id=payload.from_node_id,
        to_node_id=payload.to_node_id,
        relation=payload.relation,
        metadata_=payload.metadata_,
    )
    db.add(edge)
    await db.commit()
    await db.refresh(edge)
    return edge


@router.delete("/edges/{edge_id}", status_code=204)
async def delete_edge(
    edge_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    """Löscht eine Kante."""
    edge = await db.get(ContextEdge, edge_id)
    if not edge:
        raise HTTPException(status_code=404, detail="Kante nicht gefunden")
    
    # Prüfe Schreibrecht auf from_node
    from_node = await db.get(ContextNode, edge.from_node_id)
    if not from_node or from_node.status != "active":
        raise HTTPException(status_code=404, detail="Startknoten nicht gefunden")
    
    if from_node.write_scope == "private" and from_node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Schreibberechtigung auf Startknoten")
    if from_node.write_scope == "subject" or from_node.write_scope == "group":
        if from_node.write_scope_group_id:
            is_member = await db.execute(
                sa.select(1).where(
                    GroupMembership.group_id == from_node.write_scope_group_id,
                    GroupMembership.pseudonym == user.sub,
                )
            )
            if not is_member.scalar_one_or_none():
                raise HTTPException(status_code=403, detail="Keine Schreibberechtigung auf Startknoten")
    
    await db.delete(edge)
    await db.commit()


@router.get("/nodes/{node_id}/edges", response_model=list[ContextEdgeRead])
async def get_node_edges(
    node_id: UUID,
    relation: list[str] | None = Query(default=None),
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(get_current_user),
):
    """Gibt alle ausgehenden Kanten eines Knotens zurück."""
    # Prüfe ob Knoten existiert und sichtbar ist
    node = await db.get(ContextNode, node_id)
    if not node or node.status != "active":
        raise HTTPException(status_code=404, detail="Knoten nicht gefunden")
    
    # Sichtbarkeitsprüfung
    if node.read_scope == "private" and node.owner_pseudonym != user.sub:
        raise HTTPException(status_code=403, detail="Keine Berechtigung")
    
    query = select(ContextEdge).where(
        ContextEdge.from_node_id == node_id,
    )
    if relation:
        query = query.where(ContextEdge.relation.in_(relation))
    
    edges = (await db.execute(query)).scalars().all()
    return edges


# ── KS-Phase-6 Curriculum Create Endpoint ──────────────────────────────────


@router.post("/curricula/new", response_model=ContextNodeRead, status_code=201)
async def create_curriculum_node(
    payload: CurriculumCreate,
    db: AsyncSession = Depends(get_db),
    user: JwtPayload = Depends(_TEACHER_OR_ADMIN),
):
    """Erstellt einen neuen leeren Curriculum-Knoten für den Editor."""
    from app.db.models import Group
    
    # Subject laden
    subject = await db.execute(
        sa.select(Group.subject_id).where(
            sa.or_(
                Group.type == "subject_department",
                Group.type == "subject",
            ),
            Group.metadata_["fach_code"].astext == payload.fach_code,
        ).limit(1)
    )
    subject_row = subject.fetchone()
    if not subject_row:
        # Alternative: direkt aus subjects Tabelle
        subj = await db.execute(
            sa.select(Subject.id).where(
                Subject.metadata_["fach_code"].astext == payload.fach_code
            ).limit(1)
        )
        subject_id = subj.scalar_one_or_none()
    else:
        subject_id = subject_row[0]
    
    if not subject_id:
        raise HTTPException(
            status_code=422,
            detail=f"Fach mit fach_code '{payload.fach_code}' nicht gefunden"
        )
    
    # Fachschafts-Gruppen-ID
    department_group = await db.execute(
        sa.select(Group.id).where(
            Group.subject_id == subject_id,
            Group.type == "subject_department",
        ).limit(1)
    )
    department_group_id = department_group.scalar_one_or_none()
    
    # Prüfe ob User Mitglied der Fachschaft ist
    if department_group_id:
        is_member = await db.execute(
            sa.select(1).where(
                GroupMembership.group_id == department_group_id,
                GroupMembership.pseudonym == user.sub,
            )
        )
        if not is_member.scalar_one_or_none() and "admin" not in user.roles:
            raise HTTPException(
                status_code=403,
                detail="Keine Berechtigung - Sie müssen Mitglied der Fachschaft sein"
            )
    
    # Fachplan laden
    fachplan = await db.execute(
        sa.select(ContextNode).where(
            ContextNode.content_type == "fachplan",
            ContextNode.metadata_["fachplan_id"].astext == payload.fachplan_id,
            ContextNode.status == "active",
        ).limit(1)
    )
    fachplan_node = fachplan.scalar_one_or_none()
    
    # Curriculum-Knoten erstellen
    import_key = f"new_{user.sub}_{payload.fach_code}_{payload.jahrgangsstufe}"
    
    curriculum = ContextNode(
        category="knowledge",
        content_type="curriculum",
        title=f"{payload.fach_code} {payload.schulart} Kl. {payload.jahrgangsstufe}",
        content=None,
        read_scope="school",
        write_scope="subject",
        write_scope_group_id=department_group_id,
        subject_id=subject_id,
        owner_pseudonym=user.sub,
        metadata_={
            "fachplan_id": payload.fachplan_id,
            "bp_version": payload.bp_version,
            "schule": payload.schule,
            "fach_code": payload.fach_code,
            "schulart": payload.schulart,
            "jahrgangsstufe": payload.jahrgangsstufe,
            "import_key": import_key,
        },
        status="active",
    )
    db.add(curriculum)
    await db.commit()
    await db.refresh(curriculum)
    
    # part_of-Kante zum Fachplan
    if fachplan_node:
        edge = ContextEdge(
            from_node_id=curriculum.id,
            to_node_id=fachplan_node.id,
            relation="part_of",
            metadata_={},
        )
        db.add(edge)
        await db.commit()
    
    return curriculum


# ── KS-Phase-6 LLM-based PDF Extraction (Stufe 1) ────────────────────────────

import asyncio
import copy
import json
from dataclasses import dataclass
from typing import Literal

# Timeout für LLM-Extraktion
_CURRICULUM_TIMEOUT = 120

# Spaltennamen für bessere LLM-Orientierung
_COL_NAMES = {0: "PK", 1: "IK", 2: "Konkretisierung", 3: "Hinweise"}


@dataclass
class PageBlock:
    """Serialisierte Repräsentation einer PDF-Seite für das LLM."""
    page_number: int
    tables: list[str]  # Serialisierte Tabellen
    flow_text: str  # Fließtext der Seite


@dataclass
class ChapterChunk:
    """Ein Kapitel-Chunk mit Seitenbereichen."""
    start_page: int
    end_page: int
    pages: list[PageBlock]


# System Prompt für die LLM-Extraktion (Revision 2 - verbatim nur, keine Schlüssel)
EXTRACTION_SYSTEM_PROMPT = """Du bist ein Expert:innen-System für die Extraktion von Bildungsplan-Curricula des Landes Baden-Württemberg.

Deine Aufgabe: Analysiere die serialisierten PDF-Tabellen und extrahiere die Daten in ein strukturiertes JSON-Format.

SPALTENZUORDNUNG (STRIKT EINHALTEN):
- C0 = Prozessbezogene Kompetenzen (PK) — z.B. "2.4 Mit symbolischen Elementen umgehen"
- C1 = Inhaltsbezogene Kompetenzen (IK) — z.B. "3.1.1 Zahlbereiche", "(1) Prinzipien beschreiben"
- C2 = Konkretisierung / Vorgehen — konkrete Umsetzung im Unterricht
- C3 = Hinweise — methodische Hinweise, Materialien, Differenzierung

HINWEIS: Jede Zeile in der Serialisierung zeigt ALLE Spalten (auch leere als Cn:""); 
         die Spaltenheader "Spalten: C0:PK | C1:IK | C2:Konkretisierung | C3:Hinweise" zeigt die Bedeutung.

WICHTIGE REGELN (STRIKT einhalten):
1. SPALTENZUORDNUNG PRO ZEILE:
   - C0-Inhalt → pk_raw (Prozessbezogene Kompetenzen)
   - C1-Inhalt → ik_raw (Inhaltsbezogene Kompetenzen) ODER bp_titel (wenn Lernsequenz-Header)
   - C2-Inhalt → konkretisierung
   - C3-Inhalt → hinweise

2. Zeilentyp klassifizieren:
   - kapitel_kopf: Titel + "ca. N Std." oder "ca. N Stunden" → kapitel.titel und kapitel.std
   - lernsequenz_header: C1 enthält NUR IK-Notation wie "3.1.1 Zahlbereiche erkunden" UND C0, C2, C3 sind leer → neue Lernsequenz mit bp_titel aus C1, ik_abschnitt extrahieren
   - datenzeile: alle anderen Zeilen mit Inhalt → ein Eintrag

3. VERBATIM ÜBERNEHMEN — das ist die zentrale Regel:
   - pk_raw = C0-Inhalt KOMMPLETT und unverändert (mit ↵ als Zeilentrenner)
   - ik_raw = C1-Inhalt KOMMPLETT und unverändert (alle (N)-Items am Stück, mit ↵)
   - konkretisierung = C2-Inhalt WÖRTLICH
   - hinweise = C3-Inhalt VOLLSTÄNDIG und wörtlich (KEINE Kürzung!)
   
   WICHTIG: Auch wenn C0 oder C1 leer ist, müssen die Felder pk_raw und ik_raw als leere Strings gesetzt werden!

3. Aktuellen IK-Abschnitt mitführen:
   - Bei lernsequenz_header: ik_abschnitt = der IK-Code (z.B. "3.1.1")
   - Bei datenzeilen: ik_abschnitt der aktuellen Lernsequenz weitergeben

4. MERGED PK-Zellen:
   - Ist C0 leer/None → pk_raw = None und pk_merged_from_above = true
   - Die Vererbung übernimmt die Normalisierung, NICHT das LLM

5. KEINE REFERENZ-SCHLÜSSEL BAUEN: NICHT "3.1.1.(1)" oder "2.4.1" erzeugen!
   NICHT eckige Klammern um IK-Items entfernen!
   NICHT Texte über mehrere Einträge duplizieren!
   Das LLM liefert NUR die Roh-Daten.

6. KEIN ERFINDEN: Unklare Felder leer lassen + Warnung hinzufügen.
   Antworte NUR mit dem JSON-Objekt, kein zusätzlicher Text.

BEISPIELE:

Format A (mit Lernsequenz-Header):
=== Seite 1, Tabelle 1 (5 Zeilen × 4 Spalten) ===
Spalten: C0:PK | C1:IK | C2:Konkretisierung | C3:Hinweise
[Z0] C0:"Prozessbezogene Kompetenzen" | C1:"Inhaltsbezogene Kompetenzen" | C2:"Konkretisierung / Vorgehen" | C3:"Hinweise / Bemerkungen"
→ ÜBERSPRINGEN (Tabellenüberschriften)

[Z1] C0:"", C1:"Die Schülerinnen und Schüler können", C2:"", C3:""
→ ÜBERSPRINGEN

[Z2] C0:"", C1:"3.1.1 Zahlbereiche erkunden, Mit Zahlen Rechnen", C2:"", C3:""
→ Lernsequenz: bp_titel="3.1.1 Zahlbereiche erkunden, Mit Zahlen Rechnen", ik_abschnitt="3.1.1"

[Z3] C0:"2.5 Kommunizieren↵1. mathematische Einsichten schriftlich dokumentieren↵2.4 Mit symbolischen Elementen umgehen↵1. zwischen natürlicher und symbolischer Sprache wechseln↵3. zwischen Darstellungen wechseln↵5. Routineverfahren anwenden", C1:"(1) Prinzipien des Stellenwertsystems beschreiben↵(2) natürliche Zahlen bis Billion lesen↵(18) Zahlenwerte runden↵(6) [Zahlen und Punkte auf der Zahlengeraden zuordnen]", C2:"Natürliche Zahlen↵Große Zahlen↵Zahlen runden", C3:"Hinweis auf den Grundschulbildungsplan: „den Aufbau des Stellenwertsystems nutzen"↵Prinzipien in Analogie zum Dualsystem herausarbeiten↵MINT: Umrechnung vom Binärsystem ins Hexadezimalsystem"
→ 1 Eintrag: pk_raw="2.5 Kommunizieren↵1. …↵2.4 Mit symbolischen Elementen umgehen↵1. …", 
   ik_raw="(1) Prinzipien des Stellenwertsystems beschreiben↵(2) natürliche Zahlen bis Billion lesen↵(18) Zahlenwerte runden↵(6) [Zahlen und Punkte auf der Zahlengeraden zuordnen]",
   konkretisierung="Natürliche Zahlen↵Große Zahlen↵Zahlen runden",
   hinweise="Hinweis auf den Grundschulbildungsplan: „den Aufbau des Stellenwertsystems nutzen"↵Prinzipien in Analogie zum Dualsystem herausarbeiten↵MINT: Umrechnung vom Binärsystem ins Hexadezimalsystem"

Format B (ohne Lernsequenz-Header):
=== Seite 1, Tabelle 1 (5 Zeilen × 4 Spalten) ===
Spalten: C0:PK | C1:IK | C2:Konkretisierung | C3:Hinweise
[Z0] C0:"2.1 Argumentieren", C1:"3.1.1", C2:"Lineare Gleichungen lösen", C3:"Bezug zu den Basiskonzepten: Funktion, Algebra"
→ 1 Eintrag: ik_abschnitt="3.1.1", pk_raw="2.1 Argumentieren", ik_raw="", konkretisierung="Lineare Gleichungen lösen", hinweise="Bezug zu den Basiskonzepten: Funktion, Algebra"

Antworte STRIKT im angegebenen JSON-Schema. """


# JSON Schema für RawKapitelExtraction (LLM-Output, Revision 2)
_RAW_KAPITEL_EXTRACT_SCHEMA = {
    "name": "raw_kapitel_extract",
    "schema": {
        "type": "object",
        "properties": {
            "kapitel": {
                "type": "object",
                "properties": {
                    "titel": {"type": "string", "description": "Kapitel-Titel"},
                    "std": {"type": ["string", "null"], "description": "Stundenangabe z.B. '5 Std.'"},
                    "einleitung": {"type": ["string", "null"], "description": "Einleitungstext (Format B)"},
                    "lernsequenzen": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "bp_titel": {"type": ["string", "null"], "description": "Titel der Lernsequenz (verbatim)"},
                                "ik_abschnitt": {"type": ["string", "null"], "description": "Aktueller IK-Abschnitt z.B. '3.1.1'"},
                                "eintraege": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "ik_raw": {"type": ["string", "null"], "description": "Kompletter C1-Text, alle (N)-Items am Stück"},
                                            "pk_raw": {"type": ["string", "null"], "description": "Kompletter C0-Text, verbatim"},
                                            "pk_merged_from_above": {"type": "boolean", "default": False, "description": "True wenn PK von oben vererbt werden soll"},
                                            "konkretisierung": {"type": ["string", "null"], "description": "C2 wörtlich"},
                                            "hinweise": {"type": ["string", "null"], "description": "C3 vollständig"},
                                            "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0, "default": 1.0},
                                            "warnings": {
                                                "type": "array",
                                                "items": {"type": "string"},
                                                "default": []
                                            }
                                        },
                                        "required": ["ik_raw", "pk_raw", "konkretisierung", "hinweise"]
                                    }
                                },
                                "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0, "default": 1.0},
                                "warnings": {
                                    "type": "array",
                                    "items": {"type": "string"},
                                    "default": []
                                }
                            },
                            "required": ["bp_titel", "eintraege"]
                        }
                    },
                    "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0, "default": 1.0},
                    "warnings": {
                        "type": "array",
                        "items": {"type": "string"},
                        "default": []
                    }
                },
                "required": ["titel", "lernsequenzen"]
            }
        },
        "required": ["kapitel"]
    },
    "strict": True
}


# Altes Schema für Fallback (wird nicht mehr verwendet, bleibt für Kompatibilität)
_KAPITEL_EXTRACT_SCHEMA = {
    "name": "kapitel_extract",
    "schema": {
        "type": "object",
        "properties": {
            "kapitel": {
                "type": "object",
                "properties": {
                    "titel": {"type": "string", "description": "Kapitel-Titel"},
                    "reihenfolge": {"type": "integer", "description": "Reihenfolge-Nummer"},
                    "std": {"type": ["string", "null"], "description": "Stundenangabe"},
                    "hinweis": {"type": ["string", "null"], "description": "Einleitungstext"},
                    "konkretisierung": {"type": "array", "items": {"type": "string"}},
                    "lernsequenzen": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "bp_titel": {"type": ["string", "null"]},
                                "bp_leitidee": {"type": ["string", "null"]},
                                "reihenfolge": {"type": ["integer", "null"]},
                                "eintraege": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "ik": {"type": ["string", "null"]},
                                            "ik_partiell": {"type": "boolean", "default": False},
                                            "pk": {"type": "array", "items": {"type": "string"}},
                                            "konkretisierung": {"type": ["string", "null"]},
                                            "hinweise": {"type": ["string", "null"]},
                                            "lp": {"type": "array", "items": {"type": "string"}},
                                            "confidence": {"type": "number", "default": 1.0},
                                            "warnings": {"type": "array", "items": {"type": "string"}, "default": []}
                                        },
                                        "required": ["ik", "pk", "konkretisierung"]
                                    }
                                },
                                "confidence": {"type": "number", "default": 1.0},
                                "warnings": {"type": "array", "items": {"type": "string"}, "default": []}
                            },
                            "required": ["bp_titel", "eintraege"]
                        }
                    },
                    "confidence": {"type": "number", "default": 1.0},
                    "warnings": {"type": "array", "items": {"type": "string"}, "default": []}
                },
                "required": ["titel", "reihenfolge", "lernsequenzen"]
            }
        },
        "required": ["kapitel"]
    },
    "strict": True
}


def _strictify_schema(node: Any) -> Any:
    """Macht ein JSON-Schema OpenAI-strict-konform (in-place, rekursiv).

    OpenAIs Structured-Output-Strict-Modus verlangt für JEDES Objekt:
    - `additionalProperties: false`
    - alle Property-Keys in `required` (Optionalität nur über nullable-Typ ausdrücken)
    und lehnt Keywords wie `default`/`minimum`/`maximum` ab. Das von Hand gepflegte
    Schema bleibt für den json_object-Fallback unverändert lesbar;
    diese Funktion erzeugt daraus die strikte Variante.
    """
    if isinstance(node, dict):
        for kw in ("default", "minimum", "maximum"):
            node.pop(kw, None)
        if node.get("type") == "object" and "properties" in node:
            node["additionalProperties"] = False
            node["required"] = list(node["properties"].keys())
        for value in node.values():
            _strictify_schema(value)
    elif isinstance(node, list):
        for item in node:
            _strictify_schema(item)
    return node


_RAW_KAPITEL_EXTRACT_SCHEMA_STRICT = {
    "name": _RAW_KAPITEL_EXTRACT_SCHEMA["name"],
    "schema": _strictify_schema(copy.deepcopy(_RAW_KAPITEL_EXTRACT_SCHEMA["schema"])),
    "strict": True,
}


_KAPITEL_EXTRACT_SCHEMA_STRICT = {
    "name": _KAPITEL_EXTRACT_SCHEMA["name"],
    "schema": _strictify_schema(copy.deepcopy(_KAPITEL_EXTRACT_SCHEMA["schema"])),
    "strict": True,
}


def _extract_tables_from_pdf(content: bytes) -> list[list[list[str | None]]]:
    """Extrahiert alle Tabellenzeilen aus einem PDF via pdfplumber."""
    import pdfplumber
    rows: list[list[str | None]] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                rows.extend(table)
    return rows


def _serialize_pdf_for_llm(content: bytes) -> list[PageBlock]:
    """Serialisiert PDF-Inhalt für das LLM mit Tabellen und Fließtext.
    
    Erzeugt eine kompakte Repräsentation, die Layout-Hinweise erhält:
    - Pro Seite, pro Tabelle: ALLE Spalten mit Index (C0, C1, C2, C3)
    - Auch leere Zellen werden als Cn:"" dargestellt (wichtig für korrekte Zuordnung!)
    - Fließtext der Seite als Kontextblock
    - PDF-Silbentrennungen werden aufgelöst
    """
    import pdfplumber
    from pdfminer.high_level import extract_text
    
    pages: list[PageBlock] = []
    
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_number = page_idx + 1
            
            # Tabellen extrahieren und serialisieren
            tables_serialized: list[str] = []
            for table_idx, table in enumerate(page.extract_tables()):
                if not table or not table[0]:
                    continue
                
                num_cols = len(table[0])
                table_lines: list[str] = []
                
                for row_idx, row in enumerate(table):
                    if not row or all(c is None or c.strip() == "" for c in row):
                        continue
                    
                    row_parts: list[str] = []
                    # ALLE Spalten ausgeben (auch leere) für korrekte Zuordnung
                    for col_idx in range(num_cols):
                        cell = row[col_idx] if col_idx < len(row) else None
                        
                        # PDF-Silbentrennungen auflösen ("Lö-\nsungswege" → "Lösungswege")
                        if cell is not None:
                            cleaned = re.sub(r'-\s*\n\s*', '', cell).strip()
                        else:
                            cleaned = ""
                        
                        # Eingebettete Zeilenumbrüche als sichtbares ↵ darstellen
                        display = cleaned.replace("\n", "↵")
                        row_parts.append(f"C{col_idx}:\"{display}\"")

                    if row_parts:
                        table_lines.append(f"[Z{row_idx}] {' | '.join(row_parts)}")
                
                if table_lines:
                    header = f"=== Seite {page_number}, Tabelle {table_idx + 1} ({len(table)} Zeilen × {num_cols} Spalten) ==="
                    # Spaltenheader hinzufügen für bessere Orientierung des LLM
                    col_headers = " | ".join([f"C{col}:{_COL_NAMES.get(col, '?')}" for col in range(num_cols)])
                    tables_serialized.append(header)
                    tables_serialized.append(f"Spalten: {col_headers}")
                    tables_serialized.extend(table_lines)
            
            # Fließtext extrahieren (pdfminer nutzt 0-basierte Seitennummern)
            try:
                flow_text = extract_text(io.BytesIO(content), page_numbers=[page_idx])
            except Exception:
                flow_text = ""
            
            pages.append(PageBlock(
                page_number=page_number,
                tables=tables_serialized,
                flow_text=flow_text
            ))

    # Dokumentkopfzeilen entfernen: Zeilen, die auf ≥30% der Seiten (mind. 3) identisch
    # wiederkehren, sind Seiten-Header/-Footer und kein inhaltlicher Text.
    if len(pages) >= 3:
        from collections import Counter
        line_counts: Counter = Counter()
        for p in pages:
            seen: set[str] = set()
            for line in p.flow_text.splitlines():
                s = line.strip()
                if s and s not in seen:
                    line_counts[s] += 1
                    seen.add(s)
        min_occurrences = max(3, len(pages) // 3)
        repeated = {line for line, n in line_counts.items() if n >= min_occurrences}
        if repeated:
            pages = [
                PageBlock(
                    page_number=p.page_number,
                    tables=p.tables,
                    flow_text="\n".join(
                        line for line in p.flow_text.splitlines()
                        if line.strip() not in repeated
                    ),
                )
                for p in pages
            ]

    return pages


def _serialize_word_for_llm(content: bytes) -> list[PageBlock]:
    """Serialisiert Word-Inhalt für das LLM (analog zu PDF)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        
        pages: list[PageBlock] = []
        current_page = PageBlock(page_number=1, tables=[], flow_text="")
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text.strip():
                current_page.flow_text += text + "\n"
        
        # Tabellen extrahieren
        for table_idx, table in enumerate(doc.tables):
            table_lines: list[str] = []
            for row_idx, row in enumerate(table.rows):
                row_parts: list[str] = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_parts.append(f"C{col_idx}:\"{cell_text}\"")
                if row_parts:
                    table_lines.append(f"[Z{row_idx}] {" | ".join(row_parts)}")
            
            if table_lines:
                header = f"=== Seite 1, Tabelle {table_idx + 1} ({len(table.rows)} Zeilen × {len(table.columns)} Spalten) ==="
                current_page.tables.append(header)
                current_page.tables.extend(table_lines)
        
        pages.append(current_page)
        return pages
        
    except ImportError:
        logger.warning("python-docx nicht installiert - versuche textract")
        try:
            import textract
            text = textract.process(content).decode("utf-8", errors="replace")
            return [PageBlock(page_number=1, tables=[], flow_text=text)]
        except ImportError:
            raise ImportError("Weder python-docx noch textract installiert")


def _chunk_pages_by_kapitel(pages: list[PageBlock], max_pages_per_call: int = 4) -> list[ChapterChunk]:
    """Gruppiert Seiten in Kapitel-Chunks basierend auf Stunden-Kopfzeilen.
    
    Strategie:
    1. Kapitel-Grenzen über 'ca. N Std.' oder 'ca. N Stunden' im Fließtext erkennen
    2. Pro Kapitel ein Chunk erstellen
    3. Falls keine Kapitel gefunden: feste Fenster mit Überlappung
    """
    import re
    
    if not pages:
        return []
    
    # Kapitel-Grenzen erkennen über Stunden-Muster
    chapter_pattern = re.compile(r'ca\.\s*\d+\s*(?:Std\.?|Stunden)', re.IGNORECASE)
    
    chapter_boundaries: list[int] = [0]  # Start bei Seite 0
    
    for page_idx, page in enumerate(pages):
        # Im Fließtext nach Kapitel-Muster suchen
        if chapter_pattern.search(page.flow_text):
            # Kapitel beginnt auf dieser Seite
            if page_idx > 0 and page_idx != chapter_boundaries[-1]:
                # Vermeide Duplikate
                if page_idx not in chapter_boundaries:
                    chapter_boundaries.append(page_idx)
    
    # Immer letzte Seite als Boundary hinzufügen
    if chapter_boundaries[-1] != len(pages):
        chapter_boundaries.append(len(pages))
    
    # Falls nur eine Boundary (kein Kapitel gefunden) oder zu viele Seiten pro Chunk
    if len(chapter_boundaries) <= 2:
        # Fallback: feste Fenster mit Überlappung
        chunks = []
        page_idx = 0
        chunk_num = 0
        while page_idx < len(pages):
            end_page = min(page_idx + max_pages_per_call, len(pages))
            chunks.append(ChapterChunk(
                start_page=page_idx + 1,
                end_page=end_page,
                pages=pages[page_idx:end_page]
            ))
            # Überlappung von 1 Seite für Kontinuität
            page_idx = end_page - 1 if end_page > page_idx + 1 else end_page
            chunk_num += 1
        return chunks
    
    # Kapitel-Chunks erstellen
    chunks = []
    for i in range(len(chapter_boundaries) - 1):
        start = chapter_boundaries[i]
        end = chapter_boundaries[i + 1]
        if start < end:
            chunks.append(ChapterChunk(
                start_page=start + 1,
                end_page=end,
                pages=pages[start:end]
            ))
    
    # Letzten Chunk anpassen, falls er zu groß ist
    for chunk in chunks:
        if len(chunk.pages) > max_pages_per_call:
            # Aufteilen in kleinere Chunks
            sub_chunks = []
            page_idx = 0
            while page_idx < len(chunk.pages):
                sub_end = min(page_idx + max_pages_per_call, len(chunk.pages))
                sub_chunks.append(ChapterChunk(
                    start_page=chunk.start_page + page_idx,
                    end_page=chunk.start_page + sub_end - 1,
                    pages=chunk.pages[page_idx:sub_end]
                ))
                page_idx = sub_end
            chunks.extend(sub_chunks)
            chunks.remove(chunk)
    
    return chunks


async def _call_extraction_llm(
    serialized_content: str,
    chapter_index: int,
    settings: "Settings",
) -> dict:
    """Ruft das LLM über den LiteLLM-Proxy für die Extraktion auf.

    Primär: response_format=json_schema (Structured Output).
    Fallback: response_format=json_object + Schema im Prompt, danach tolerantes Parsen.
    """
    import httpx
    from app.config import settings as app_settings

    model = settings.curriculum_extract_model or app_settings.chat_default_model

    base_payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
            {"role": "user", "content": serialized_content},
        ],
        "stream": False,
        "temperature": 0,
        "user": "curriculum-import",
    }

    # Vom Proxy gemeldeter Upstream-Fehler mit erhaltenem HTTP-Status, damit der
    # Aufrufer einen 400 (Parameter nicht unterstützt) von echten Fehlern (4xx/5xx)
    # unterscheiden kann.
    class _LLMUpstreamError(Exception):
        def __init__(self, status_code: int, text: str):
            self.status_code = status_code
            self.text = text
            super().__init__(f"{status_code}: {text}")

    async def _post(payload: dict) -> str:
        async with httpx.AsyncClient(timeout=_CURRICULUM_TIMEOUT) as client:
            resp = await client.post(
                f"{app_settings.litellm_proxy_url}/chat/completions",
                headers={"Authorization": f"Bearer {app_settings.litellm_master_key}"},
                json=payload,
            )
        if resp.status_code == 429:
            raise HTTPException(status_code=429, detail="Budget erschöpft - bitte später erneut versuchen")
        if resp.status_code != 200:
            raise _LLMUpstreamError(resp.status_code, resp.text)
        data = resp.json()
        if not data.get("choices"):
            raise ValueError(f"Unerwartetes Antwortformat: {data}")
        return data["choices"][0]["message"]["content"]

    # Gestufte Strategie: json_schema → json_object → ganz ohne response_format.
    # Ein 400 (Modell unterstützt den response_format-Modus nicht) oder unparsebares
    # JSON degradiert auf die nächste Stufe; 429/Verbindung/sonstige Status brechen ab.
    schema_hint = json.dumps(_RAW_KAPITEL_EXTRACT_SCHEMA["schema"], ensure_ascii=False, indent=2)
    fallback_messages = [
        {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT + f"\n\nAntworte STRIKT im folgenden JSON-Schema:\n{schema_hint}"},
        {"role": "user", "content": serialized_content},
    ]
    attempts = [
        ("json_schema", {**base_payload, "response_format": {"type": "json_schema", "json_schema": _RAW_KAPITEL_EXTRACT_SCHEMA_STRICT}}),
        ("json_object", {**base_payload, "messages": fallback_messages, "response_format": {"type": "json_object"}}),
        ("plain", {**base_payload, "messages": fallback_messages}),
    ]

    try:
        last_err: Exception | None = None
        for mode, payload in attempts:
            try:
                raw = await _post(payload)
                # _parse_llm_response strippt Markdown-Fences und parst tolerant.
                return _parse_llm_response(raw)
            except _LLMUpstreamError as err:
                if err.status_code != 400:
                    # Echter Proxy-/Provider-Fehler (z. B. 401/404/5xx) → nicht degradieren.
                    raise HTTPException(status_code=502, detail=f"LLM-Request fehlgeschlagen: {err.text}")
                last_err = err
                logger.warning(
                    f"Modus '{mode}' vom Modell abgelehnt (Kapitel {chapter_index}): "
                    f"{err.text[:200]} — nächste Stufe"
                )
            except (json.JSONDecodeError, ValueError) as err:
                last_err = err
                logger.warning(
                    f"Modus '{mode}' lieferte ungültiges JSON (Kapitel {chapter_index}): {err} — nächste Stufe"
                )
        # Alle Stufen erschöpft → Aufrufer erzeugt Platzhalter-Kapitel mit Warnung.
        raise HTTPException(status_code=502, detail=f"LLM-Request fehlgeschlagen (alle Modi): {last_err}")

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail=f"LLM-Timeout nach {_CURRICULUM_TIMEOUT}s")
    except httpx.ConnectError:
        raise HTTPException(status_code=502, detail="Keine Verbindung zum LiteLLM-Proxy")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"LLM-Aufruf fehlgeschlagen (Kapitel {chapter_index}): {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _parse_llm_response(content: str) -> dict[str, Any]:
    """Parsed die LLM-Antwort falls kein json_schema unterstützt wird.
    
    Versuch JSON zu extrahieren, auch wenn es in Markdown-Fences eingebettet ist.
    """
    # Markdown-Fences entfernen
    content = re.sub(r'^```(?:json)?\s*', '', content, flags=re.MULTILINE)
    content = re.sub(r'\s*```\s*$', '', content, flags=re.MULTILINE)
    content = content.strip()
    
    # Wirft bei unparsebarem Inhalt — der Aufrufer (_call_extraction_llm) degradiert dann
    # auf die nächste Stufe bzw. erzeugt am Ende ein Platzhalter-Kapitel mit Warnung.
    return json.loads(content)


async def _extract_curriculum_via_llm(
    content: bytes,
    fachplan_id: str,
    bp_version: str,
    fach: str = "",
    jahrgangsstufe: str = "",
    schulart: str = "",
) -> "CurriculumDraftData":
    """Extrahiert Curriculum-Struktur via LLM (Revision 2).
    
    1. Seiten serialisieren
    2. In Kapitel-Chunks gruppieren
    3. Pro Chunk LLM-Call → RawKapitelExtraction (verbatim)
    4. Normalisierung: RawKapitelExtraction → CurriculumDraftKapitel
    5. Ergebnisse zusammenführen
    """
    from app.config import settings as app_settings
    from app.context.schemas import CurriculumDraftData, RawKapitelExtraction
    from app.context.curriculum_normalize import normalize_raw_extraction
    
    # Serialisierung
    try:
        pages = _serialize_pdf_for_llm(content)
    except Exception as e:
        logger.error(f"PDF-Serialisierung fehlgeschlagen: {e}")
        raise HTTPException(status_code=500, detail=f"PDF-Serialisierung fehlgeschlagen: {str(e)}")
    
    if not pages:
        raise HTTPException(status_code=400, detail="Keine Seiten im PDF gefunden")
    
    # Chunking
    max_pages = app_settings.curriculum_extract_max_pages_per_call
    chunks = _chunk_pages_by_kapitel(pages, max_pages)
    
    if not chunks:
        raise HTTPException(status_code=400, detail="Keine Kapitel im Dokument erkannt")
    
    # Parallel LLM-Calls ausführen (begrenzt via Semaphore)
    concurrency = app_settings.curriculum_extract_concurrency
    semaphore = asyncio.Semaphore(concurrency)
    
    async def process_chunk(chunk: ChapterChunk, index: int) -> dict[str, Any]:
        """Verarbeitet einen Chunk und gibt RawKapitelExtraction-Rohdaten zurück."""
        async with semaphore:
            # Chunk für LLM serialisieren
            chunk_text_parts = []
            for page in chunk.pages:
                if page.tables:
                    chunk_text_parts.extend(page.tables)
                if page.flow_text.strip():
                    chunk_text_parts.append(f"--- Fließtext Seite {page.page_number} ---")
                    chunk_text_parts.append(page.flow_text.strip())
            
            serialized_input = "\n".join(chunk_text_parts)
            
            # LLM aufrufen (mit Retry bei Validierungsfehler)
            for attempt in range(2):  # 1 Original + 1 Retry
                try:
                    result = await _call_extraction_llm(serialized_input, index, app_settings)
                    
                    # Rohdaten validieren
                    raw_chapter = RawKapitelExtraction.model_validate(result.get("kapitel", {}))
                    return {"index": index, "raw": result, "validated": raw_chapter}
                    
                except Exception as e:
                    if attempt == 0:
                        logger.warning(f"Erster Versuch fehlgeschlagen (Kapitel {index}): {e}, Retry...")
                        # Warte kurz vor Retry
                        await asyncio.sleep(1)
                        continue
                    else:
                        logger.error(f"LLM-Extraktion fehlgeschlagen (Kapitel {index}): {e}")
                        # Platzhalter-Rohdaten
                        return {
                            "index": index,
                            "raw": {
                                "kapitel": {
                                    "titel": f"Kapitel {index + 1} (Fehler bei Extraktion)",
                                    "lernsequenzen": []
                                }
                            },
                            "validated": RawKapitelExtraction(
                                titel=f"Kapitel {index + 1} (Fehler bei Extraktion)",
                                lernsequenzen=[],
                                warnings=[f"Extraktion fehlgeschlagen: {str(e)}"]
                            )
                        }
    
    # Chunks verarbeiten
    raw_results = await asyncio.gather(*[process_chunk(chunk, index) for index, chunk in enumerate(chunks)])
    
    # Normalisierung: Raw → CurriculumDraftKapitel
    normalized_chapters = []
    for raw_result in raw_results:
        try:
            # Normalisierung durchführen
            chapters = normalize_raw_extraction(raw_result["raw"])
            normalized_chapters.extend(chapters)
        except Exception as e:
            logger.error(f"Normalisierung fehlgeschlagen: {e}")
            # Fallback: Platzhalter-Kapitel
            normalized_chapters.append({
                "titel": f"Kapitel (Normalisierungsfehler)",
                "reihenfolge": len(normalized_chapters) + 1,
                "lernsequenzen": [],
                "warnings": [f"Normalisierung fehlgeschlagen: {str(e)}"]
            })
    
    # Zusammenführen zu CurriculumDraftData
    return CurriculumDraftData(
        schule="",
        fach_code=fach[:2].upper() if fach else "??",
        fach=fach or None,
        schulart=schulart,
        jahrgangsstufe=jahrgangsstufe,
        fachplan_id=fachplan_id,
        bp_version=bp_version,
        vorwort=None,
        kapitel=normalized_chapters,
    )


def _extract_text_from_pdf(content: bytes) -> str:
    """Extrahiert Text aus PDF-Datei (via pdfminer.six)."""
    from pdfminer.high_level import extract_text
    return extract_text(io.BytesIO(content))


def _detect_format(text_content: str) -> tuple[bool, str | None, list[str]]:
    """Erkennt ob das Dokument unterstützt wird.
    
    Rückgabe: (is_supported, format_detected, warnings)
    - is_supported: True wenn Curriculum, False wenn nicht unterstützt (z.B. Kompetenzmatrix)
    - format_detected: "A", "B" oder None (nur für UI-Badge)
    - warnings: Liste mit Warnungen
    """
    warnings = []
    text_lower = text_content.lower()
    
    # Abbruch: Kompetenzmatrix-Dokument
    if "kompetenzmatrix" in text_lower:
        return False, None, ["Kompetenzmatrix-Format erkannt – nicht unterstützt"]
    
    # Pflicht-Signale für ein BW-Schulcurriculum
    has_pk_header = bool(re.search(r'prozessbezogene\s+kompetenzen', text_lower))
    has_ik_header = bool(re.search(r'inhaltsbezogene\s+kompetenzen', text_lower))
    
    # IK-Nummern: dreistellig (3.1.1) oder zweistellig (3.1)
    ik_numbers = re.findall(r'\b\d+\.\d+(?:\.\d+)?\b', text_content)
    # PK-Nummern: zweistellig am Zeilenanfang (2.5 Kommunizieren)
    pk_numbers = re.findall(r'(?m)^\s*\d+\.\d+\s+[A-ZÄÖÜ]', text_content)
    
    has_curriculum_structure = (
        (has_pk_header or has_ik_header)
        and len(ik_numbers) >= 3
    ) or (len(ik_numbers) >= 5 and len(pk_numbers) >= 3)
    
    if not has_curriculum_structure:
        return False, None, ["Kein erkennbares Curriculum-Format gefunden"]
    
    # Format erkennung für UI-Badge
    has_lernsequenz = bool(re.search(r'\blernsequenz', text_lower))
    format_detected = "A" if has_lernsequenz else "B"
    
    return True, format_detected, warnings


def _extract_text_from_word(content: bytes) -> str:
    """Extrahiert Text aus Word-Dokument."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join([p.text for p in doc.paragraphs])
    except ImportError:
        logger.warning("python-docx nicht installiert - versuche textract")
        try:
            import textract
            return textract.process(content).decode("utf-8", errors="replace")
        except ImportError:
            raise ImportError("Weder python-docx noch textract installiert")


def _extract_lp_codes(text: str) -> list[str]:
    """Extrahiert Leitperspektive-Codes aus Text."""
    patterns = [
        r'\bL\s+[A-Z]{2,3}\b',  # L BO
        r'\(L\)\s+[A-Z]{2,4}',  # (L) BTV
        r'\bLP[-_]?[A-Z0-9]+\b',  # LP-XX
    ]
    codes = []
    for pattern in patterns:
        codes.extend(re.findall(pattern, text, re.IGNORECASE))
    return codes


