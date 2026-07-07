"""Export-Funktionen für Stundenentwürfe: Markdown, PDF, DOCX."""

from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from uuid import UUID

import sqlalchemy as sa
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models import ContextEdge, ContextNode, LessonSlot

logger = logging.getLogger(__name__)

_TEMPLATES_DIR = Path(__file__).parent / "templates"


@dataclass
class ExportPhase:
    name: str
    dauer_min: int
    beschreibung: str
    prio: str
    sozialform: str  # display string
    methode: str  # display string
    material: list[str]  # display strings


@dataclass
class LessonExport:
    titel: str
    titel_slug: str
    stundenziel: str
    gruppe: str
    gruppe_slug: str
    datum: str  # YYYY-MM-DD or ""
    start_period: Optional[int]
    periods: int
    verfuegbare_min: int
    ue_titel: str
    phasen: list[ExportPhase]
    refs: list[dict]  # [{typ, code, titel, partiell}]


def _slugify(s: str) -> str:
    s = s.lower().strip()
    s = re.sub(r"[äöü]", lambda m: {"ä": "ae", "ö": "oe", "ü": "ue"}[m.group()], s)
    s = re.sub(r"[^a-z0-9]+", "-", s)
    return s.strip("-")[:60] or "stunde"


def _display_linked(item: dict) -> str:
    if not item:
        return ""
    if item.get("typ") == "node":
        return item.get("titel") or item.get("node_id", "")
    return item.get("wert") or ""


async def build_lesson_export(db: AsyncSession, node_id: UUID) -> LessonExport:
    lesson = await db.get(ContextNode, node_id)
    if not lesson:
        raise ValueError(f"Stunde {node_id} nicht gefunden")

    meta = lesson.metadata_ or {}

    # Slot
    slot_result = await db.execute(
        sa.select(LessonSlot).where(LessonSlot.stunde_node_id == node_id)
    )
    slot = slot_result.scalar_one_or_none()

    # UE
    ue_edge = (
        await db.execute(
            sa.select(ContextEdge).where(
                ContextEdge.from_node_id == node_id,
                ContextEdge.relation == "part_of",
            )
        )
    ).scalar_one_or_none()
    ue_titel = ""
    if ue_edge:
        ue_node = await db.get(ContextNode, ue_edge.to_node_id)
        if ue_node:
            ue_titel = ue_node.title

    # Gruppe
    gruppe = ""
    if lesson.write_scope_group_id:
        from app.db.models import Group
        g = await db.get(Group, lesson.write_scope_group_id)
        if g:
            gruppe = g.name

    datum = slot.date.isoformat() if slot else ""
    periods = slot.periods if slot else 1
    start_period = slot.start_period if slot else None

    phasen = [
        ExportPhase(
            name=p.get("name", ""),
            dauer_min=p.get("dauer_min", 0),
            beschreibung=p.get("beschreibung") or "",
            prio=p.get("prio", "kern"),
            sozialform=_display_linked(p.get("sozialform") or {}),
            methode=_display_linked(p.get("methode") or {}),
            material=[_display_linked(m) for m in (p.get("material") or [])],
        )
        for p in meta.get("phasen", [])
    ]

    return LessonExport(
        titel=lesson.title,
        titel_slug=_slugify(lesson.title),
        stundenziel=meta.get("stundenziel") or "",
        gruppe=gruppe,
        gruppe_slug=_slugify(gruppe) if gruppe else "gruppe",
        datum=datum,
        start_period=start_period,
        periods=periods,
        verfuegbare_min=periods * 45,
        ue_titel=ue_titel,
        phasen=phasen,
        refs=meta.get("refs", []),
    )


def export_markdown(data: LessonExport) -> str:
    prio_labels = {"kern": "Kern", "uebung": "Übung", "vertiefung": "Vertiefung"}
    lines = [
        "---",
        f"titel: {data.titel}",
        f"datum: {data.datum}",
        f"gruppe: {data.gruppe}",
        f"ue: {data.ue_titel}",
        f"verfuegbar_min: {data.verfuegbare_min}",
    ]
    if data.refs:
        kompetenzen = [
            f"{r.get('code', r.get('titel', ''))}{'[…]' if r.get('partiell') else ''}"
            for r in data.refs
        ]
        lines.append(f"kompetenzen: [{', '.join(kompetenzen)}]")
    lines += ["---", ""]

    if data.stundenziel:
        lines += [f"**Stundenziel:** {data.stundenziel}", ""]

    total = sum(p.dauer_min for p in data.phasen)
    lines.append(
        f"Zeitbudget: {total}′ geplant / {data.verfuegbare_min}′ verfügbar"
        + (f" (**+{total - data.verfuegbare_min}′ Überhang**)" if total > data.verfuegbare_min else "")
    )
    lines.append("")

    for phase in data.phasen:
        prio = prio_labels.get(phase.prio, phase.prio)
        lines.append(f"## {phase.name} ({phase.dauer_min}′ · {prio})")
        if phase.beschreibung:
            lines.append("")
            lines.append(phase.beschreibung)
        if phase.sozialform:
            lines.append(f"- **Sozialform:** {phase.sozialform}")
        if phase.methode:
            lines.append(f"- **Methode:** {phase.methode}")
        if phase.material:
            for m in phase.material:
                if m:
                    lines.append(f"- **Material:** {m}")
        lines.append("")

    return "\n".join(lines)


async def export_pdf(data: LessonExport) -> bytes:
    try:
        import weasyprint
        from jinja2 import Environment, FileSystemLoader, select_autoescape
    except ImportError as e:
        raise RuntimeError(
            f"PDF-Export benötigt 'weasyprint' und 'jinja2' ({e})."
        ) from e

    # Phase 17 (D5): Beschreibungen/Stundenziel als Inline-Markdown mit server-gerenderten
    # Formeln ($…$ → eingebettetes SVG) vor-rendern (kompakte Tabellenzellen → inline).
    from app.render.export import render_markdown_inline_for_pdf

    beschreibungen = [await render_markdown_inline_for_pdf(p.beschreibung or "") for p in data.phasen]
    stundenziel_html = await render_markdown_inline_for_pdf(data.stundenziel or "")

    env = Environment(
        loader=FileSystemLoader(str(_TEMPLATES_DIR)),
        autoescape=select_autoescape(["html"]),
    )
    template = env.get_template("lesson.html")
    html_str = template.render(
        data=data, beschreibungen=beschreibungen, stundenziel_html=stundenziel_html
    )
    return weasyprint.HTML(string=html_str, base_url=str(_TEMPLATES_DIR)).write_pdf()


def export_docx(data: LessonExport) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError(f"DOCX-Export benötigt 'python-docx' ({e}).") from e

    import io

    prio_labels = {"kern": "Kern", "uebung": "Übung", "vertiefung": "Vertiefung"}
    prio_abbr = {"kern": "K", "uebung": "Ü", "vertiefung": "V"}
    prio_hex = {"kern": "4A7FB5", "uebung": "5BA37A", "vertiefung": "B07FB8"}

    def _shade_cell(cell, hex_color: str) -> None:
        """Hintergrundfarbe einer Tabellenzelle setzen (python-docx kann das nicht direkt)."""
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Kopf
    h = doc.add_heading(data.titel, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_parts = []
    if data.datum:
        meta_parts.append(data.datum)
    if data.gruppe:
        meta_parts.append(data.gruppe)
    if data.ue_titel:
        meta_parts.append(f"UE: {data.ue_titel}")
    if meta_parts:
        doc.add_paragraph(" · ".join(meta_parts)).runs[0].font.size = Pt(10)

    if data.stundenziel:
        p = doc.add_paragraph()
        p.add_run("Stundenziel: ").bold = True
        p.add_run(data.stundenziel)

    total = sum(ph.dauer_min for ph in data.phasen)
    budget_text = f"Zeitbudget: {total}′ / {data.verfuegbare_min}′ verfügbar"
    if total > data.verfuegbare_min:
        budget_text += f" (+{total - data.verfuegbare_min}′ Überhang)"
    doc.add_paragraph(budget_text).runs[0].font.size = Pt(9)

    doc.add_paragraph()

    if data.phasen:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Zeit", "Prio", "Phase", "Sozialform / Methode", "Material"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True

        kum = 0
        for phase in data.phasen:
            row = table.add_row().cells
            row[0].text = f"{kum}–{kum + phase.dauer_min}′"

            # Prio: farbige Zelle + Kürzel (K/Ü/V), weiß zentriert — analog zur App.
            prio_cell = row[1]
            prio_cell.text = prio_abbr.get(phase.prio, "?")
            _shade_cell(prio_cell, prio_hex.get(phase.prio, "888888"))
            p = prio_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            row[2].text = f"{phase.name}\n{phase.beschreibung}" if phase.beschreibung else phase.name

            # Sozialform (fett) über Methode stapeln.
            mc = row[3]
            mc.text = phase.sozialform or phase.methode or ""
            if phase.sozialform:
                mc.paragraphs[0].runs[0].bold = True
                if phase.methode:
                    mc.add_paragraph(phase.methode)

            row[4].text = "\n".join(m for m in phase.material if m)
            kum += phase.dauer_min

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
